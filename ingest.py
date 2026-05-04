from dotenv import load_dotenv
load_dotenv()

import os
import json
import logging
import re
import time
from collections import Counter

import chromadb
from chromadb.utils import embedding_functions

from docx import Document
from pypdf import PdfReader
from openai import OpenAI

import pdfplumber
import tiktoken

# ==========================
# ENV
# ==========================
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not OPENAI_API_KEY:
    raise ValueError("Falta OPENAI_API_KEY")

os.environ["CHROMA_OPENAI_API_KEY"] = OPENAI_API_KEY
os.environ["ANONYMIZED_TELEMETRY"] = "False"

# ==========================
# LOGGING
# ==========================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s"
)
logger = logging.getLogger(__name__)

# ==========================
# PATHS
# ==========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CHROMA_PATH = os.path.join(BASE_DIR, "chroma_db")

# 🔴 NUEVO: logging en disco
LOG_DIR = os.path.join(BASE_DIR, "logs", "chunks")
os.makedirs(LOG_DIR, exist_ok=True)

# ==========================
# CLIENTS
# ==========================
client = chromadb.PersistentClient(path=CHROMA_PATH)

embedding_function = embedding_functions.OpenAIEmbeddingFunction(
    api_key=OPENAI_API_KEY,
    model_name="text-embedding-3-large"
)

collection = client.get_or_create_collection(
    name="documents",
    embedding_function=embedding_function
)

llm = OpenAI(api_key=OPENAI_API_KEY)

# ==========================
# CLEAN (PRO)
# ==========================
def clean(text):
    lines = text.split("\n")

    lines = [l.strip() for l in lines if l.strip() and len(l.strip()) > 3]
    counts = Counter(lines)

    cleaned = []
    for l in lines:
        low = l.lower()

        if counts[l] > 3:
            continue
        if re.match(r"^page \d+", low):
            continue
        if re.match(r"^\d+$", l):
            continue
        if re.match(r"^\d+\.\s+[a-z]", low):
            continue
        if "...." in l:
            continue
        if "confidential" in low:
            continue

        cleaned.append(l)

    return "\n".join(cleaned)

# ==========================
# SAFE FILENAME (NUEVO)
# ==========================
import unicodedata

def safe_filename(name):
    # quitar acentos
    name = unicodedata.normalize('NFKD', name).encode('ascii', 'ignore').decode('ascii')

    # eliminar caracteres inválidos en Windows
    name = re.sub(r'[<>:"/\\|?*]', '', name)

    # reemplazar espacios por _
    name = re.sub(r'\s+', '_', name)

    # quitar espacios inicio/fin
    name = name.strip()

    return name

# ==========================
# LOG CHUNKS
# ==========================
def log_chunks(doc_name, chunks):
    safe_name = safe_filename(doc_name)
    folder = os.path.join(LOG_DIR, safe_name)
    os.makedirs(folder, exist_ok=True)

    path = os.path.join(folder, "chunks.jsonl")

    with open(path, "w", encoding="utf-8") as f:
        for i, c in enumerate(chunks):
            f.write(json.dumps({
                "chunk_index": i,
                "length": len(c),
                "chunk": c
            }, ensure_ascii=False) + "\n")

    logger.info(f"[LOG] Chunks guardados en: {folder}")


# ==========================
# PARSE PDF
# ==========================
def parse_pdf(path):
    logger.info(f"Parsing PDF: {path}")
    logger.info("[TRACE] Inicio parse PDF")

    text_blocks = []
    text_found = False

    try:
        reader = PdfReader(path)
        for p in reader.pages:
            content = p.extract_text()
            if content and len(content.strip()) > 50:
                text_found = True
                text_blocks.append(content)
    except Exception as e:
        logger.warning(f"Error leyendo PDF: {e}")

    try:
        with pdfplumber.open(path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                tables = page.extract_tables()

                for table in tables:
                    for row in table:
                        row_text = " | ".join([str(cell) for cell in row if cell])
                        if row_text.strip():
                            text_blocks.append(row_text)

    except Exception as e:
        logger.warning(f"Error tablas: {e}")

    if not text_found:
        logger.warning(f"PDF sin texto extraíble (posible escaneado): {path}")
        return ""

    logger.info(f"[TRACE] Fin parse PDF | bloques: {len(text_blocks)}")

    full_text = "\n".join(text_blocks)
    return clean(full_text)

# ==========================
# PARSE DOCX
# ==========================
def parse_docx(path):
    doc = Document(path)

    text_blocks = []

    for p in doc.paragraphs:
        if p.text.strip():
            text_blocks.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                text_blocks.append(row_text)

    return clean("\n".join(text_blocks))

# ==========================
# CHUNK
# ==========================
def chunk(text, max_tokens=500, overlap=80):
    enc = tiktoken.encoding_for_model("gpt-4.1-mini")
    tokens = enc.encode(text)

    chunks = []

    i = 0
    while i < len(tokens):
        chunk_tokens = tokens[i:i + max_tokens]
        chunk_text = enc.decode(chunk_tokens)

        if chunk_text.strip():
            chunks.append(chunk_text.strip())

        i += max_tokens - overlap

    return chunks

# ==========================
# INGEST
# ==========================
def ingest(path):
    start_time = time.time()

    logger.info(f"Ingest: {path}")

    raw_name = os.path.splitext(os.path.basename(path))[0]
    name = safe_filename(raw_name)

    if path.endswith(".pdf"):
        text = parse_pdf(path)
    else:
        text = parse_docx(path)

    if not text:
        logger.warning("Documento vacío")
        return

    # 🔴 control documento grande
    MAX_CHARS = 200000
    if len(text) > MAX_CHARS:
        logger.warning(f"[LARGE DOC] {name} truncado")
        text = text[:MAX_CHARS]

    chunks = chunk(text)

    # 🔴 logging chunks
    log_chunks(name, chunks)
    logger.info(f"{name}: {len(chunks)} chunks generados")


    # 🔴 batch embeddings
    BATCH_SIZE = 100

    doc_id = name  # 🔥 clave para trazabilidad
    # 🔴 evitar duplicados
    try:
        existing = collection.get(where={"doc_id": doc_id})
        if existing["ids"]:
           logger.warning(f"[SKIP] Documento ya indexado: {doc_id}")
           return
    except Exception as e:
        logger.warning(f"[CHECK ERROR] {e}")

    BATCH_SIZE = 50

    for i in range(0, len(chunks), BATCH_SIZE):

      batch = chunks[i:i+BATCH_SIZE]

      summary_docs = []
      summary_meta = []
      summary_ids = []

      full_docs = []
      full_meta = []
      full_ids = []

      for j, ch in enumerate(batch):
          idx = i + j

          # 🔴 SUMMARY (puedes mejorar luego con LLM si quieres)
          summary = ch[:300] if len(ch) > 300 else ch

          summary_docs.append(summary)
          summary_meta.append({
              "type": "summary",
              "doc_id": doc_id,
              "chunk_index": idx
          })
          summary_ids.append(f"{doc_id}_summary_{idx}")

          # 🔴 FULL
          full_docs.append(ch)
          full_meta.append({
              "type": "full",
              "doc_id": doc_id,
              "chunk_index": idx
          })
          full_ids.append(f"{doc_id}_full_{idx}")

      for attempt in range(3):
          try:
            collection.add(
              ids=summary_ids,
              documents=summary_docs,
              metadatas=summary_meta
        )

            collection.add(
              ids=full_ids,
              documents=full_docs,
              metadatas=full_meta
        )

            logger.info(f"[BATCH] Insertados {len(summary_ids)} summary + {len(full_ids)} full")
            break

          except Exception as e:
             logger.error(f"[RETRY {attempt}] {e}")
             time.sleep(1)
             if attempt == 2:
               logger.error("[FAIL] Batch descartado")
        

    # 🔴 métricas
    duration = round(time.time() - start_time, 2)
    coverage = round(len(text) / MAX_CHARS, 2) if len(text) >= MAX_CHARS else 1

    logger.info(f"""
[METRICS]
Documento: {name}
Chunks: {len(chunks)}
Chars: {len(text)}
Tiempo: {duration}s
Coverage: {coverage}
""")

    logger.info(f"✔ Ingest completado: {name}")

# ==========================
# QUERY EMBEDDING
# ==========================
def embed_query(query):
    return embedding_function([query])[0]

# ==========================
# SEARCH (SUMMARY FIRST)
# ==========================
def search_summary(query, n_results=20):
    results = collection.query(
        query_texts=[query],
        n_results=n_results,
        where={"type": "summary"}
    )
    return results

# ==========================
# FETCH FULL CHUNKS
# ==========================
def fetch_full_chunks(summary_results):
    full_ids = []

    for meta in summary_results.get("metadatas", [[]])[0]:
        doc_id = meta.get("doc_id")
        idx = meta.get("chunk_index") 
        if not doc_id or idx is None:
             continue

        full_ids.append(f"{doc_id}_full_{idx}")

        if not full_ids:
            return []

    try:
        full_docs = collection.get(ids=full_ids)
    except Exception as e:
        logger.error(f"[ERROR fetch_full] {e}")
        return []

    return full_docs["documents"]
# ==========================
# DEDUPE
# ==========================
def dedupe_chunks(chunks):
    seen = set()
    unique = []

    for ch in chunks:
        key = ch[:200]  # hash simple

        if key not in seen:
            seen.add(key)
            unique.append(ch)

    return unique
# ==========================
# EMBEDDING RERANK (NEW)
# ==========================
def embedding_rerank(query, chunks):
    """
    Rerank semántico usando embeddings + cosine similarity
    """

    if not chunks:
        return []

    try:
        # 1. embedding query
        query_emb = embedding_function([query])[0]

        # 2. embedding chunks
        chunk_embs = embedding_function(chunks)

        # 3. cosine similarity
        def cosine(a, b):
            dot = sum(x*y for x, y in zip(a, b))
            norm_a = sum(x*x for x in a) ** 0.5
            norm_b = sum(x*x for x in b) ** 0.5
            return dot / (norm_a * norm_b + 1e-8)

        scored = []

        for ch, emb in zip(chunks, chunk_embs):
            score = cosine(query_emb, emb)
            scored.append((score, ch))

        # ordenar descendente
        scored.sort(key=lambda x: x[0], reverse=True)
        for score, ch in scored[:5]:
            logger.info(f"[RERANK SCORE] {round(score, 3)} | {ch[:80]}")

        return [ch for _, ch in scored]

    except Exception as e:
        logger.error(f"[RERANK ERROR] {e}")
        return chunks  # fallback seguro

# ==========================
# RETRIEVAL PIPELINE
# ==========================
def retrieve(query, top_k=5):

    # 1. buscar summaries
    summary_results = search_summary(query, n_results=30)

    # 2. traer full chunks
    full_chunks = fetch_full_chunks(summary_results)

    # 3. dedupe
    unique_chunks = dedupe_chunks(full_chunks)
    MAX_RERANK = 25
    unique_chunks = unique_chunks[:MAX_RERANK]

    # 4. rerank
    ranked_chunks = embedding_rerank(query, unique_chunks)
    # 5. top_k final
    final_chunks = ranked_chunks[:top_k]

    return final_chunks
# ==========================
# BUILD CONTEXT
# ==========================
def build_context(chunks, max_chars=4000):
    context = ""

    for ch in chunks:
        if len(context) + len(ch) > max_chars:
            break
        context += ch + "\n\n"

    return context
# ==========================
# ANSWER
# ==========================
def answer(query):

    chunks = retrieve(query)
    context = build_context(chunks)

    prompt = f"""
Usa el siguiente contexto para responder la pregunta.

Contexto:
{context}

Pregunta:
{query}
"""

    response = llm.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "user", "content": prompt}
        ]
    )

    return response.choices[0].message.content


# ==========================
# CLI
# ==========================
if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser()

    # modos
    parser.add_argument("--ingest", type=str, help="Carpeta con documentos")
    parser.add_argument("--query", type=str, help="Consulta al sistema RAG")

    args = parser.parse_args()

    # ==========================
    # INGEST
    # ==========================
    if args.ingest:
        logger.info(f"Carpeta ingest: {args.ingest}")

        for f in os.listdir(args.ingest):
            if f.endswith((".pdf", ".docx")):
                ingest(os.path.join(args.ingest, f))

    # ==========================
    # QUERY
    # ==========================
    if args.query:
        logger.info(f"Query: {args.query}")

        result = answer(args.query)

        print("\n==================== RESPUESTA ====================\n")
        print(result)
        print("\n===================================================\n")